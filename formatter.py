from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, font_name="Lexend", size_px=16):
    """Add a heading with default paragraph spacing."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_px * 0.75)
    return doc


def add_list_item(doc, text, font_name="Lexend Light", size_px=12):
    """Add a bulleted list item with default spacing."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_px * 0.75)
    return doc


def add_formatted_paragraph(
    doc, placeholder, user_text,
    placeholder_font="Lexend", text_font="Lexend Light",
    size_px=12
):
    """Add a paragraph with differently formatted placeholder and user text."""
    p = doc.add_paragraph()

    run_placeholder = p.add_run(f"{placeholder}: ")
    run_placeholder.font.name = placeholder_font
    run_placeholder.font.size = Pt(size_px * 0.75)

    run_text = p.add_run(user_text)
    run_text.font.name = text_font
    run_text.font.size = Pt(size_px * 0.75)

    return doc


def add_horizontal_line(doc):
    """Add a thin horizontal line with default paragraph spacing."""
    p = doc.add_paragraph()

    p_border = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')         # thin line
    bottom_border.set(qn('w:color'), '000000') # black
    p_border.append(bottom_border)
    p._p.get_or_add_pPr().append(p_border)

    return doc


def add_subheading_bold(doc, left_text, right_text, font_name="Lexend", size_px=12):
    """Add subheading 1 as a two-column table with default spacing."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Left cell
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run(left_text)
    left_run.font.name = font_name
    left_run.font.size = Pt(size_px * 0.75)
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Right cell
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_run = right_para.add_run(right_text)
    right_run.font.name = font_name
    right_run.font.size = Pt(size_px * 0.75)
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def add_subheading_italic(doc, left_text, right_text, font_name="Lexend Light", size_px=12, italic=True):
    """Add subheading 2 as a two-column table with default spacing."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Left cell
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run(left_text)
    left_run.font.name = font_name
    left_run.font.size = Pt(size_px * 0.75)
    left_run.font.italic = italic
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Right cell
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_run = right_para.add_run(right_text)
    right_run.font.name = font_name
    right_run.font.size = Pt(size_px * 0.75)
    right_run.font.italic = italic
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def add_contact_info(doc, email, phone, location, font_name="Lexend", size_px=12):
    """Add a single-line contact info, left-aligned, with default spacing."""
    p = doc.add_paragraph()
    text = f"{email} | {phone} | {location}"
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_px * 0.75)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return doc