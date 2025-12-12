from docx import Document
from formatter import (
    add_heading, add_list_item, add_formatted_paragraph,
    add_horizontal_line, add_subheading_bold, add_subheading_italic,
    add_contact_info
)

def clean_ai(text: str):
    """Remove <<start>> and <<end>> markers if present."""
    if not text:
        return ""
    return text.replace("<<start>>", "").replace("<<end>>", "").strip()


def build_resume(data, output_file="resume.docx"):
    doc = Document()

    # PERSONAL INFO
    personal = data["personalInfo"]
    add_heading(doc, personal.get("fullName", ""))
    add_horizontal_line(doc)
    add_contact_info(
        doc,
        personal.get("email", ""),
        personal.get("phone", ""),
        personal.get("location", "")
    )

    # WORK EXPERIENCE
    work_experiences = data.get("workExperiences", [])
    if work_experiences:
        add_heading(doc, "WORK EXPERIENCE")
        add_horizontal_line(doc)

        for exp in work_experiences:
            # Combine join/exit dates
            join = exp.get("joinDate", "")
            exit_ = exp.get("exitDate", "")
            dates_combined = f"{join} — {exit_}" if join or exit_ else ""

            add_subheading_bold(doc, exp.get("company", ""), dates_combined)
            add_subheading_italic(doc, exp.get("role", ""), exp.get("office", ""))

            # Clean AI markers
            description = clean_ai(exp.get("description", ""))
            add_list_item(doc, description)

            responsibility = clean_ai(exp.get("responsibility", ""))
            add_list_item(doc, responsibility)

            achievements = clean_ai(exp.get("achievements", ""))
            add_list_item(doc, achievements)

            tools = clean_ai(exp.get("tools", ""))
            add_list_item(doc, tools)

            projects = clean_ai(exp.get("projects", ""))
            add_list_item(doc, projects)

    # EDUCATION
    education_history = data.get("educationHistory", [])
    if education_history:
        add_heading(doc, "EDUCATION")
        add_horizontal_line(doc)

        for edu in education_history:
            add_subheading_bold(doc, edu.get("instituteName", ""), edu.get("instituteLocation", ""))
            add_subheading_italic(doc, edu.get("degreeTitle", ""), edu.get("graduationYear", ""))
            add_list_item(doc, f"Major: {edu.get('majorSubject', '')}")
            add_list_item(doc, f"Minor: {edu.get('minorSubject', '')}")
            highlight = edu.get("highlight", "")
            if highlight:
                add_list_item(doc, f"Highlight: {highlight}")

    # SKILLS & INTERESTS
    add_heading(doc, "SKILLS & INTERESTS")
    add_horizontal_line(doc)

    technologies = data.get("technology", "")
    if technologies:
        add_formatted_paragraph(doc, "Technologies", technologies.rstrip(".") + ".")

    soft_skills = data.get("softSkills", "")
    if soft_skills:
        add_formatted_paragraph(doc, "Soft Skills", soft_skills.rstrip(".") + ".")

    interests = data.get("interest", "")
    if interests:
        add_formatted_paragraph(doc, "Interests", interests.rstrip(".") + ".")

    # CERTIFICATIONS
    certifications = data.get("certifications", [])
    if certifications:
        add_heading(doc, "CERTIFICATIONS")
        add_horizontal_line(doc)

        for cert in certifications:
            add_subheading_bold(
                doc,
                cert.get("title", ""),
                f"{cert.get('issuer', '')}, {cert.get('date', '')}"
            )
            # Correct key name: 'courses'
            for course in cert.get("courses", []):
                add_list_item(doc, course)

    doc.save(output_file)
    return output_file
